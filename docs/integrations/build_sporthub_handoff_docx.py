from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "NUTRIO_SPORTHUB_DEVELOPER_HANDOFF.md"
OUTPUT = ROOT / "Nutrio-SportHub-Developer-Handoff.docx"

INK = RGBColor(0x02, 0x06, 0x17)
GREEN = RGBColor(0x22, 0xC7, 0xA1)
BLUE = RGBColor(0x38, 0xBD, 0xF8)
PURPLE = RGBColor(0x7C, 0x83, 0xF6)
MUTED = RGBColor(0x64, 0x74, 0x8B)
LIGHT = "F6F8FB"


def set_font(run, name="Arial", size=11, color=INK, bold=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2

for style_name, size, color, before, after in [
    ("Heading 1", 17, INK, 16, 8),
    ("Heading 2", 13, GREEN, 13, 6),
    ("Heading 3", 11.5, PURPLE, 10, 4),
]:
    style = doc.styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run("NUTRIO x SPORTHUB  |  INTEGRATION CONTRACT 1.0")
set_font(run, size=8.5, color=MUTED, bold=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Confidential partner engineering handoff  |  ")
set_font(run, size=8, color=MUTED)
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

# Customer-pack opening block.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(26)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("PARTNER ENGINEERING HANDOFF")
set_font(r, size=10, color=GREEN, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Nutrio x SportHub")
set_font(r, size=28, color=INK, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("OAuth, activity API, signed webhooks, data rules, and joint acceptance tests")
set_font(r, size=13, color=MUTED)

meta = doc.add_table(rows=2, cols=3)
meta.autofit = False
widths = [2.2, 2.2, 2.2]
values = [
    ("CONTRACT", "Version 1.0"), ("STATUS", "Ready for sandbox"), ("DATE", "12 July 2026"),
    ("OWNER", "Nutrio Engineering"), ("REGION", "Qatar"), ("DATA MODE", "Minimum necessary"),
]
for index, (label, value) in enumerate(values):
    cell = meta.cell(index // 3, index % 3)
    cell.width = Inches(widths[index % 3])
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label + "\n")
    set_font(r, size=7.5, color=GREEN, bold=True)
    r = p.add_run(value)
    set_font(r, size=9.5, color=INK, bold=True)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    tc_pr.append(shd)

doc.add_paragraph()

lines = SOURCE.read_text(encoding="utf-8").splitlines()
in_code = False
code_lines = []
skip_title = True

def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=9, color=PURPLE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=10.5, color=INK, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=10.5, color=INK)


for line in lines:
    if line.startswith("```"):
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            shade(p, LIGHT)
            run = p.add_run("\n".join(code_lines))
            set_font(run, name="Consolas", size=8.4, color=INK)
            code_lines = []
            in_code = False
        else:
            in_code = True
        continue
    if in_code:
        code_lines.append(line)
        continue
    if skip_title and line.startswith("# "):
        skip_title = False
        continue
    if line.startswith("### "):
        doc.add_paragraph(line[4:], style="Heading 3")
    elif line.startswith("## "):
        doc.add_paragraph(line[3:], style="Heading 2")
    elif line.startswith("# "):
        doc.add_paragraph(line[2:], style="Heading 1")
    elif re.match(r"^\d+\. ", line):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        add_inline(p, re.sub(r"^\d+\. ", "", line))
    elif line.startswith("- [ ] "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.24)
        add_inline(p, "☐ " + line[6:])
    elif line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        add_inline(p, line[2:])
    elif line.startswith("**") and line.endswith("  "):
        p = doc.add_paragraph()
        add_inline(p, line.strip())
    elif line.strip():
        p = doc.add_paragraph()
        add_inline(p, line)

doc.save(OUTPUT)
print(OUTPUT)
